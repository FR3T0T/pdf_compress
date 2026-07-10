"""
pdf_translate — Offline translation & OCR engine.

Translates PDF text and text inside images (photos, scans) entirely
on-device:

  * Translation  — Argos Translate (CTranslate2 / OpenNMT models).  Runs
                   fully offline once the language packages are installed.
  * OCR          — Tesseract (via pytesseract) for pulling text out of
                   images and scanned pages, with per-language trained data.
  * Detection    — langdetect for source-language auto-detection.

Nothing here makes a network request.  The language models and OCR data
are provisioned separately and explicitly by ``setup_translation.py`` (the
one online step); after that, everything runs locally.

The module is import-safe: heavy optional libraries (argostranslate,
pytesseract, langdetect, PyMuPDF, python-docx) are imported lazily so the
rest of the toolkit keeps working even when translation isn't provisioned
yet.  Functions that need a missing piece raise :class:`TranslationError`
with a clear, actionable message.
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass
from typing import Callable, Optional

import translate_runtime

log = logging.getLogger(__name__)

MAX_FILE_SIZE = 2 * 1024 * 1024 * 1024  # 2 GB, matches the rest of the toolkit


# ════════════════════════════════════════════════════════════════════
#  Errors
# ════════════════════════════════════════════════════════════════════

class TranslationError(Exception):
    """Base error for translation/OCR problems."""


class ModelMissingError(TranslationError):
    """A required language model or OCR pack is not installed."""


# ════════════════════════════════════════════════════════════════════
#  Supported languages
#
#  The genuine global top-10 by total speakers, plus German and Danish
#  (both explicitly requested).  Each entry maps our stable code to the
#  Argos translation code and the Tesseract OCR code.
# ════════════════════════════════════════════════════════════════════

@dataclass(frozen=True)
class Language:
    code: str        # our stable id (ISO 639-1 where possible)
    name: str        # English display name
    native: str      # endonym
    argos: str       # Argos Translate language code
    tesseract: str   # Tesseract OCR language code


SUPPORTED_LANGUAGES: list[Language] = [
    Language("en", "English",            "English",     "en", "eng"),
    Language("zh", "Chinese (Mandarin)", "中文",         "zh", "chi_sim"),
    Language("hi", "Hindi",              "हिन्दी",        "hi", "hin"),
    Language("es", "Spanish",            "Español",     "es", "spa"),
    Language("ar", "Arabic",             "العربية",      "ar", "ara"),
    Language("fr", "French",             "Français",    "fr", "fra"),
    Language("bn", "Bengali",            "বাংলা",        "bn", "ben"),
    Language("pt", "Portuguese",         "Português",   "pt", "por"),
    Language("ru", "Russian",            "Русский",     "ru", "rus"),
    Language("id", "Indonesian",         "Indonesia",   "id", "ind"),
    Language("de", "German",             "Deutsch",     "de", "deu"),
    Language("da", "Danish",             "Dansk",       "da", "dan"),
]

LANG_BY_CODE: dict[str, Language] = {l.code: l for l in SUPPORTED_LANGUAGES}
_ARGOS_TO_CODE = {l.argos: l.code for l in SUPPORTED_LANGUAGES}

# langdetect emits a few codes that need normalising to ours.
_LANGDETECT_FIXUP = {
    "zh-cn": "zh", "zh-tw": "zh", "zh": "zh",
}


def supported_languages() -> list[dict]:
    """Return the supported languages as plain dicts (for the UI)."""
    return [
        {"code": l.code, "name": l.name, "native": l.native}
        for l in SUPPORTED_LANGUAGES
    ]


# ════════════════════════════════════════════════════════════════════
#  Lazy backends
# ════════════════════════════════════════════════════════════════════

def _argos():
    # In a frozen build the ML stack is deliberately NOT bundled (frozen
    # PyTorch aborts natively; see translate_runtime's module docstring) —
    # activate() makes a runtime provisioned by the in-app setup flow
    # importable first. No-op in a source checkout with argos installed.
    translate_runtime.activate()
    try:
        import argostranslate.package as p
        import argostranslate.translate as t
        return t, p
    except Exception as exc:  # not installed (or the stack failed to load)
        # Surface the underlying cause: a genuine not-installed reads as
        # "No module named 'argostranslate'", while an environment problem
        # (e.g. a missing module in a frozen build) would otherwise hide
        # behind this generic message and be undebuggable from the toast.
        raise ModelMissingError(
            "Offline translation isn't set up yet. Use the Translate tool's "
            "one-time setup (downloads the translation engine and your "
            "chosen languages), or in a source checkout:\n"
            "    pip install argostranslate\n"
            "    python setup_translation.py --install all\n"
            f"({exc.__class__.__name__}: {exc})"
        ) from exc


def _pytesseract():
    try:
        import pytesseract
        return pytesseract
    except Exception as exc:
        raise ModelMissingError(
            "OCR isn't available — pytesseract is not installed "
            "(pip install pytesseract) and Tesseract must be on the system."
        ) from exc


# ════════════════════════════════════════════════════════════════════
#  Provisioning status
# ════════════════════════════════════════════════════════════════════

def _installed_ocr_langs() -> list[str]:
    """Tesseract OCR language codes available on this machine."""
    try:
        pt = _pytesseract()
        return sorted(pt.get_languages(config=""))
    except Exception:
        return []


def _installed_argos_pairs() -> list[tuple[str, str]]:
    """Installed Argos (from_code, to_code) pairs."""
    try:
        t, p = _argos()
        return sorted((pkg.from_code, pkg.to_code)
                      for pkg in p.get_installed_packages())
    except Exception:
        return []


def translation_status() -> dict:
    """Report what is provisioned so the UI can guide the user.

    Returns target languages reachable from English (the Argos pivot) and
    which OCR packs are present — never raises.
    """
    ocr = set(_installed_ocr_langs())
    pairs = set(_installed_argos_pairs())

    langs = []
    for l in SUPPORTED_LANGUAGES:
        to_ok = l.argos == "en" or ("en", l.argos) in pairs
        from_ok = l.argos == "en" or (l.argos, "en") in pairs
        langs.append({
            "code": l.code,
            "name": l.name,
            "native": l.native,
            "ocr": l.tesseract in ocr,        # can read this language from images
            "translateTo": to_ok,             # can translate INTO this language
            "translateFrom": from_ok,         # can translate FROM this language
        })

    translate_runtime.activate()
    try:
        import argostranslate  # noqa: F401
        argos_available = True
    except Exception:
        argos_available = False

    return {
        "argosAvailable": argos_available,
        "ocrAvailable": bool(ocr),
        "ocrLangs": sorted(ocr),
        "argosPairs": [f"{a}->{b}" for a, b in sorted(pairs)],
        "languages": langs,
        # Drives the in-app setup flow (frozen builds provision the ML
        # stack on demand — see translate_runtime).
        "runtime": translate_runtime.runtime_status(),
    }


# ════════════════════════════════════════════════════════════════════
#  Language pack installation (network — explicit, user-initiated only)
# ════════════════════════════════════════════════════════════════════

def install_languages(codes: list[str],
                      progress: Optional[Callable[[int, int, str], None]] = None,
                      should_cancel: Optional[Callable[[], bool]] = None) -> dict:
    """Download + install the Argos en<->X packages for ``codes``.

    THE sanctioned in-app network operation (with translate_runtime's
    wheel download): runs only when the user explicitly starts translation
    setup from the Translate tool, or via setup_translation.py's CLI.
    Argos stores the packages in its own per-user data dir, so they
    survive app upgrades and work offline afterwards.

    Returns {"installed": n, "skipped": n, "requested": n}.
    """
    _t, package = _argos()

    if "all" in codes:
        codes = [lang.code for lang in SUPPORTED_LANGUAGES]
    unknown = [c for c in codes if c not in LANG_BY_CODE]
    if unknown:
        raise TranslationError(f"Unknown language code(s): {', '.join(unknown)}")

    pairs: set[tuple[str, str]] = set()
    for c in codes:
        argos = LANG_BY_CODE[c].argos
        if argos == "en":
            continue
        pairs.add(("en", argos))
        pairs.add((argos, "en"))
    if not pairs:
        return {"installed": 0, "skipped": 0, "requested": 0}

    try:
        package.update_package_index()
        available = package.get_available_packages()
    except Exception as exc:
        raise TranslationError(
            f"Could not reach the Argos package index: {exc}") from exc

    ordered = sorted(pairs)
    ok, skipped = 0, 0
    for i, (from_code, to_code) in enumerate(ordered):
        if should_cancel and should_cancel():
            raise InterruptedError("Cancelled")
        if progress:
            progress(i, len(ordered), f"Downloading {from_code}→{to_code}")
        match = next((p for p in available
                      if p.from_code == from_code and p.to_code == to_code), None)
        if match is None:
            log.warning("no Argos package published for %s->%s", from_code, to_code)
            skipped += 1
            continue
        try:
            path = match.download()
            try:
                match.install()                      # newer Argos
            except Exception:
                package.install_from_path(path)      # older Argos
            ok += 1
        except Exception as exc:
            log.warning("installing %s->%s failed: %s", from_code, to_code, exc)
            skipped += 1
    if progress:
        progress(len(ordered), len(ordered), "Language packs installed")
    return {"installed": ok, "skipped": skipped, "requested": len(ordered)}


# ════════════════════════════════════════════════════════════════════
#  Language detection
# ════════════════════════════════════════════════════════════════════

def detect_language(text: str) -> Optional[str]:
    """Best-effort source-language detection → our code (or None)."""
    sample = (text or "").strip()
    if len(sample) < 3:
        return None
    try:
        from langdetect import DetectorFactory, detect
        DetectorFactory.seed = 0  # deterministic
        raw = detect(sample)
    except Exception:
        return None
    raw = _LANGDETECT_FIXUP.get(raw, raw)
    return raw if raw in LANG_BY_CODE else None


# ════════════════════════════════════════════════════════════════════
#  Text translation
# ════════════════════════════════════════════════════════════════════

def _resolve_source(text: str, source: str) -> str:
    if source and source != "auto":
        if source not in LANG_BY_CODE:
            raise TranslationError(f"Unsupported source language: {source}")
        return source
    detected = detect_language(text)
    if not detected:
        raise TranslationError(
            "Could not auto-detect the source language — please select it.")
    return detected


# ════════════════════════════════════════════════════════════════════
#  Protecting literals during translation
#
#  Sent raw, Argos hallucinates on lone symbols (a bare "|" becomes
#  "124"/"Yep 124") and translates proper nouns it shouldn't ("Commerce"
#  -> "Handel"). translate_line() splits a line on separator characters,
#  masks emails/URLs/"City, ST"/phone numbers/acronyms/numbers/caller-
#  supplied terms with sentinel tokens before handing the remainder to
#  the model, then restores them afterward -- the model never sees the
#  literal, so it can't mistranslate or hallucinate on it. See _sentinel()
#  for why the tokens are plain ASCII, not the originally-planned
#  private-use-area codepoints.
# ════════════════════════════════════════════════════════════════════

_SEP_RE = re.compile(r'(\s*[|•·‣▪◦]\s*|\s{2,})')
_PATTERNS = [
    re.compile(r'\b[\w.+-]+@[\w-]+\.[\w.-]+\b'),
    re.compile(r'\b(?:https?://|www\.)\S+|\b[\w-]+\.(?:com|org|net|edu|io|gov)(?:/\S*)?\b', re.I),
    re.compile(r'\b[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*,\s*[A-Z]{2}\b'),
    re.compile(r'\+?\d[\d\-\(\)\.\s]{6,}\d'),
    re.compile(r'\b(?:[A-Z]\.){1,}[A-Z]?\b'),
    re.compile(r'\b[A-Z]{2,5}\b'),
    re.compile(r'\b\d[\d,.\-/]*\b'),
]


def _sentinel(i):
    # ASCII, not the private-use-area codepoints this was originally
    # written with: verified against the installed Argos model that PUA
    # sentinels get corrupted mid-translation (e.g. \uE010 comes back as
    # a real Cyrillic letter), silently destroying the protected span on
    # restore. This single-token ASCII form survives Argos translation
    # intact in testing for the realistic case (a handful of protected
    # spans per line); very dense lines with many sentinels are not fully
    # reliable -- a limitation of asking a generative NMT model to pass
    # through unfamiliar tokens verbatim, not something a smarter
    # sentinel format alone can fully solve.
    return f"Zqpt{i}qzX"


def _protect(text, extra_terms):
    spans = []
    for term in (extra_terms or []):
        for m in re.finditer(rf'\b{re.escape(term)}\b', text):
            spans.append((m.start(), m.end()))
    for rx in _PATTERNS:
        for m in rx.finditer(text):
            spans.append((m.start(), m.end()))
    spans.sort(key=lambda s: (s[0], -(s[1] - s[0])))
    chosen, last = [], -1
    for a, b in spans:
        if a >= last:
            chosen.append((a, b))
            last = b
    saved = {}
    for i, (a, b) in enumerate(reversed(chosen)):
        key = _sentinel(i)
        saved[key] = text[a:b]
        text = text[:a] + key + text[b:]
    return text, saved


def _restore(text, saved):
    for key, val in saved.items():
        text = text.replace(key, val)
    return text


def translate_line(line, translate_fn, protect_terms=None):
    """Translate one line, protecting separators/emails/phones/acronyms/
    numbers/caller-supplied terms from the model. See module docstring
    above this section for why.
    """
    if not line.strip():
        return line
    out = []
    for part in _SEP_RE.split(line):
        if not part or _SEP_RE.fullmatch(part):
            out.append(part)
            continue
        masked, saved = _protect(part, protect_terms)
        residual = masked
        for k in saved:
            residual = residual.replace(k, "")
        # Skip a fragment only when it holds no letter in ANY script -- i.e.
        # it's purely digits/punctuation/whitespace/separators. str.isalpha()
        # is Unicode-aware, so real words in CJK, Arabic, Devanagari, Bengali,
        # etc. get translated too, not just Latin/Cyrillic (TRN-03).
        if not any(ch.isalpha() for ch in residual):
            out.append(part)
            continue
        out.append(_restore(translate_fn(masked), saved))
    return ''.join(out)


def translate_text(text: str, target: str, source: str = "auto",
                    protect_terms: Optional[list[str]] = None) -> dict:
    """Translate a block of text. Returns {translated, source, target}.

    Argos pivots through English automatically when a direct model pair
    isn't installed (e.g. de→en→ru). Routed line-by-line through
    translate_line() so separators/emails/phone numbers/acronyms/proper
    nouns like "City, ST" survive instead of being hallucinated or
    mistranslated (see the section above). ``protect_terms`` lets callers
    add their own words (names, places) the heuristic patterns wouldn't
    otherwise catch.
    """
    if not text or not text.strip():
        return {"translated": "", "source": source, "target": target}
    if target not in LANG_BY_CODE:
        raise TranslationError(f"Unsupported target language: {target}")

    src = _resolve_source(text, source)
    if src == target:
        return {"translated": text, "source": src, "target": target}

    t, _p = _argos()
    from_code = LANG_BY_CODE[src].argos
    to_code = LANG_BY_CODE[target].argos

    def _translate_fn(s: str) -> str:
        try:
            return t.translate(s, from_code, to_code)
        except Exception as exc:
            raise ModelMissingError(
                f"No offline model for {LANG_BY_CODE[src].name} → "
                f"{LANG_BY_CODE[target].name}. Install it with:\n"
                f"    python setup_translation.py --install {src} {target}"
            ) from exc

    translated = "\n".join(
        translate_line(line, _translate_fn, protect_terms)
        for line in text.split("\n")
    )
    return {"translated": translated, "source": src, "target": target}


# ════════════════════════════════════════════════════════════════════
#  OCR (image / photo text)
# ════════════════════════════════════════════════════════════════════

def ocr_image(image_path: str, source: str = "auto") -> dict:
    """Pull text out of an image with Tesseract. Returns {text, ocrLang}.

    ``source`` selects the OCR trained-data pack; "auto" tries every
    installed pack the toolkit knows about (less accurate than naming it).
    """
    if not os.path.isfile(image_path):
        raise FileNotFoundError(image_path)
    pt = _pytesseract()
    from PIL import Image

    installed = set(_installed_ocr_langs())
    if not installed:
        raise ModelMissingError(
            "No Tesseract OCR language data is installed. "
            "See setup_translation.py for the install commands.")

    if source and source != "auto":
        if source not in LANG_BY_CODE:
            raise TranslationError(f"Unsupported source language: {source}")
        ocr_code = LANG_BY_CODE[source].tesseract
        if ocr_code not in installed:
            raise ModelMissingError(
                f"OCR data for {LANG_BY_CODE[source].name} ({ocr_code}) "
                f"is not installed. See setup_translation.py.")
    else:
        # Use every known installed pack at once (Tesseract supports '+').
        known = [l.tesseract for l in SUPPORTED_LANGUAGES if l.tesseract in installed]
        ocr_code = "+".join(known) if known else "eng"

    try:
        with Image.open(image_path) as img:
            text = pt.image_to_string(img, lang=ocr_code)
    except Exception as exc:
        raise TranslationError(f"OCR failed: {exc}") from exc

    return {"text": text.strip(), "ocrLang": ocr_code}


def translate_image(image_path: str, target: str, source: str = "auto",
                     protect_terms: Optional[list[str]] = None) -> dict:
    """OCR an image then translate the extracted text.

    Returns {sourceText, translatedText, source, target, ocrLang}.
    """
    ocr = ocr_image(image_path, source)
    src_text = ocr["text"]
    if not src_text:
        return {
            "sourceText": "", "translatedText": "",
            "source": source, "target": target, "ocrLang": ocr["ocrLang"],
            "note": "No text was found in the image.",
        }
    result = translate_text(src_text, target, source, protect_terms)
    return {
        "sourceText": src_text,
        "translatedText": result["translated"],
        "source": result["source"],
        "target": target,
        "ocrLang": ocr["ocrLang"],
    }


# ════════════════════════════════════════════════════════════════════
#  PDF translation
# ════════════════════════════════════════════════════════════════════

def _extract_pages(path: str, ocr_fallback_source: str = "auto") -> list[str]:
    """Return the text of each page. OCRs pages that have no text layer."""
    import fitz  # PyMuPDF
    pages: list[str] = []
    doc = fitz.open(path)
    try:
        for page in doc:
            txt = page.get_text("text").strip()
            if not txt:
                # Scanned page — rasterize and OCR it.
                try:
                    import tempfile
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    fd, tmp = tempfile.mkstemp(suffix=".png")
                    os.close(fd)
                    # pix.save() is inside the try so the finally always cleans
                    # up the mkstemp'd file: if the save itself raises (disk
                    # full, PyMuPDF error) the temp .png still exists and would
                    # otherwise leak (TRN-02). mkstemp/os.close stay outside so
                    # `tmp` is defined for the finally.
                    try:
                        pix.save(tmp)
                        txt = ocr_image(tmp, ocr_fallback_source)["text"]
                    finally:
                        if os.path.exists(tmp):
                            os.unlink(tmp)
                except ModelMissingError:
                    raise
                except Exception as exc:
                    log.debug("page OCR fallback failed: %s", exc)
                    txt = ""
            pages.append(txt)
    finally:
        doc.close()
    return pages


# Scripts that render correctly with plain left-to-right textbox insertion
# (no shaping/reordering needed) once a font with the right coverage is
# used. zh is handled separately via fitz's built-in CJK support.
_LATIN_CYRILLIC_TARGETS = {"en", "es", "fr", "pt", "de", "da", "id", "ru"}

# PyMuPDF's insert_textbox() lays out left-to-right, one glyph after another,
# with no bidi reordering or contextual shaping. That's wrong for these
# scripts (Arabic needs RTL + joined letterforms, Hindi/Bengali need
# grapheme-cluster shaping for conjuncts) -- inserting translated text
# directly would silently produce unreadable output. PDF output for these
# targets is routed to .docx instead (see translate_pdf).
_COMPLEX_SHAPING_TARGETS = {"ar", "hi", "bn"}

_FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets", "fonts")

# A Unicode TTF covering Latin + Greek + Cyrillic, needed to embed into the
# translated PDF for _LATIN_CYRILLIC_TARGETS (fitz's built-in "helv" is
# Latin-only). Checked in order; the first one found is embedded. Bundling
# a font (e.g. DejaVuSans.ttf, SIL/Bitstream Vera license) into
# assets/fonts/ is preferred since it's portable across machines -- these
# are the OS-provided fallbacks used until one is added there. Referenced
# by absolute path at runtime only; never copied into the repo (most are
# proprietary OS assets, not redistributable).
_UNICODE_FONT_CANDIDATES = [
    os.path.join(_FONTS_DIR, "DejaVuSans.ttf"),
    os.path.join(_FONTS_DIR, "NotoSans-Regular.ttf"),
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\tahoma.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
]


def _resolve_unicode_font_path() -> Optional[str]:
    for path in _UNICODE_FONT_CANDIDATES:
        if os.path.isfile(path):
            return path
    return None


def _font_kwargs_for_target(target: str) -> dict:
    """insert_textbox() font kwargs for a translation target.

    zh uses fitz's built-in CJK support (the "china-s" reserved fontname --
    no file needed, MuPDF has it embedded). Latin/Cyrillic targets get a
    Unicode TTF if one can be found (see _resolve_unicode_font_path);
    otherwise this falls back to the Latin-only built-in font and logs a
    loud warning rather than failing silently.
    """
    if target == "zh":
        return {"fontname": "china-s"}
    if target in _LATIN_CYRILLIC_TARGETS:
        font_path = _resolve_unicode_font_path()
        if font_path:
            return {"fontfile": font_path, "fontname": "unicode-body"}
        log.warning(
            "No Unicode TTF found for PDF translation target '%s' (checked "
            "%s); falling back to the built-in Latin-only font -- Cyrillic/"
            "Greek characters will not render correctly. Drop a font (e.g. "
            "DejaVuSans.ttf) into %s to fix this.",
            target, _UNICODE_FONT_CANDIDATES, _FONTS_DIR,
        )
    return {"fontname": "helv"}


def _insert_autofit_text(page, rect, text: str, font_kwargs: dict) -> None:
    """Insert ``text`` into ``rect``, shrinking the font until it fits.

    insert_textbox() returns negative when text overflows the rect. Starts
    near a normal body size and steps down to a sane minimum -- dense
    blocks (a short source string that translates much longer) may still
    end up visually smaller than the original; that's an accepted
    trade-off of reusing the original block's rect rather than reflowing
    the whole page.
    """
    size = 11.0
    min_size = 5.0
    while size >= min_size:
        overflow = page.insert_textbox(rect, text, fontsize=size, **font_kwargs)
        if overflow >= 0:
            return
        size -= 0.5
    page.insert_textbox(rect, text, fontsize=min_size, **font_kwargs)


def _translate_pdf_to_pdf(input_path: str, output_path: str, target: str,
                           source: str,
                           progress: Optional[Callable[[int, int], None]],
                           should_cancel: Optional[Callable[[], bool]],
                           protect_terms: Optional[list[str]] = None) -> dict:
    """Rebuild the PDF page-by-page, keeping original images in place and
    replacing each text block's text with its translation in the same rect.

    Layout is approximate, not a true reflow: images keep their exact
    original position/size, but translated text rarely matches the source
    string's length, so text is auto-shrunk to fit its original block's
    rect (see _insert_autofit_text) rather than reflowing the page. Callers
    must not route _COMPLEX_SHAPING_TARGETS here -- see that constant.
    """
    import fitz  # PyMuPDF

    src = fitz.open(input_path)
    out = fitz.open()
    total = src.page_count
    detected_source = None
    font_kwargs = _font_kwargs_for_target(target)

    try:
        for i in range(total):
            if should_cancel and should_cancel():
                raise InterruptedError("Cancelled")
            page = src[i]
            new_page = out.new_page(width=page.rect.width, height=page.rect.height)

            for img in page.get_images(full=True):
                xref = img[0]
                try:
                    rects = page.get_image_rects(xref)
                    if not rects:
                        continue
                    img_bytes = src.extract_image(xref)["image"]
                    for rect in rects:
                        new_page.insert_image(rect, stream=img_bytes)
                except Exception as exc:
                    log.debug("skipping image xref=%s on page %d: %s", xref, i, exc)

            for block in page.get_text("blocks"):
                x0, y0, x1, y1, text, _block_no, block_type = block[:7]
                if block_type != 0 or not text.strip():
                    continue
                # Once the document language is known from an earlier block,
                # reuse it so short blocks (page numbers, years, <3 chars) that
                # would fail auto-detection still translate instead of raising.
                effective_source = (
                    detected_source if source == "auto" and detected_source else source
                )
                try:
                    res = translate_text(text, target, effective_source, protect_terms)
                    detected_source = detected_source or res["source"]
                    translated = res["translated"]
                except Exception as exc:
                    # An undetectable/untranslatable block (e.g. a bare number
                    # that fails detect_language) must NOT abort the document --
                    # keep it verbatim so the loop continues and out.save() is
                    # still reached (TRN-01). Covers TranslationError and any
                    # backstop failure.
                    log.debug("keeping block untranslated on page %d: %s", i, exc)
                    translated = text
                rect = fitz.Rect(x0, y0, x1, y1)
                try:
                    _insert_autofit_text(new_page, rect, translated, font_kwargs)
                except Exception as exc:
                    log.debug("skipping text block on page %d: %s", i, exc)

            if progress:
                try:
                    progress(i + 1, total)
                except Exception:
                    pass

        out.save(output_path)
    finally:
        src.close()
        out.close()

    return {
        "output": output_path,
        "outputSize": os.path.getsize(output_path),
        "pages": total,
        "source": detected_source or source,
        "target": target,
        "files": [output_path],
        "output_dir": os.path.dirname(os.path.abspath(output_path)),
    }


def translate_pdf(input_path: str, output_path: str, target: str,
                  source: str = "auto",
                  progress: Optional[Callable[[int, int], None]] = None,
                  should_cancel: Optional[Callable[[], bool]] = None,
                  protect_terms: Optional[list[str]] = None) -> dict:
    """Translate a PDF's text and write the result.

    Output format follows the ``output_path`` extension:
      * .pdf  — rebuilds each page, keeping original images in place and
                replacing text in its original block position (approximate
                layout — see _translate_pdf_to_pdf). Not available for
                _COMPLEX_SHAPING_TARGETS (ar, hi, bn); those are
                automatically redirected to .docx, with a `note` in the
                returned dict explaining why, since PyMuPDF's text
                insertion doesn't perform the bidi/shaping those scripts
                need and would otherwise produce unreadable output.
      * .docx — a Word document with a heading per page (needs python-docx)
      * .txt  — page-delimited plain text (fallback for any other/missing
                extension)

    .txt/.docx render every script correctly using the system's own fonts
    and involve no layout guesswork; .pdf is best-effort and intended for
    cases where keeping the original images and rough page layout matters
    more than pixel-perfect text placement. Returns a summary dict.
    """
    if not os.path.isfile(input_path):
        raise FileNotFoundError(input_path)
    size = os.path.getsize(input_path)
    if size > MAX_FILE_SIZE:
        raise TranslationError("File too large (max 2 GB).")
    if target not in LANG_BY_CODE:
        raise TranslationError(f"Unsupported target language: {target}")

    ext = os.path.splitext(output_path)[1].lower()
    note = None

    if ext == ".pdf" and target in _COMPLEX_SHAPING_TARGETS:
        output_path = os.path.splitext(output_path)[0] + ".docx"
        ext = ".docx"
        note = (
            f"{LANG_BY_CODE[target].name} text needs shaping that PDF output "
            f"doesn't support correctly, so this was saved as a Word "
            f"document instead."
        )

    if ext == ".pdf":
        result = _translate_pdf_to_pdf(input_path, output_path, target, source,
                                        progress, should_cancel, protect_terms)
        if note:
            result["note"] = note
        return result

    page_texts = _extract_pages(input_path, source)
    total = len(page_texts)
    translated_pages: list[str] = []
    detected_source = None

    for i, ptext in enumerate(page_texts):
        if should_cancel and should_cancel():
            raise InterruptedError("Cancelled")
        if not ptext.strip():
            translated_pages.append("")
        else:
            res = translate_text(ptext, target, source, protect_terms)
            detected_source = detected_source or res["source"]
            translated_pages.append(res["translated"])
        if progress:
            try:
                progress(i + 1, total)
            except Exception:
                pass

    if ext == ".docx":
        _write_docx(translated_pages, output_path)
    else:
        _write_txt(translated_pages, output_path)

    result = {
        "output": output_path,
        "outputSize": os.path.getsize(output_path),
        "pages": total,
        "source": detected_source or source,
        "target": target,
        "files": [output_path],
        "output_dir": os.path.dirname(os.path.abspath(output_path)),
    }
    if note:
        result["note"] = note
    return result


def _write_txt(pages: list[str], output_path: str) -> None:
    import tempfile
    out_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    fd, tmp = tempfile.mkstemp(suffix=".txt", dir=out_dir)
    os.close(fd)
    try:
        with open(tmp, "w", encoding="utf-8") as fh:
            for i, txt in enumerate(pages):
                fh.write(f"--- Page {i + 1} ---\n")
                fh.write(txt + "\n\n")
        os.replace(tmp, output_path)
    except Exception:
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise


def _write_docx(pages: list[str], output_path: str) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise TranslationError(
            "Word output needs python-docx (pip install python-docx). "
            "Use a .txt output path instead.") from exc
    doc = Document()
    for i, txt in enumerate(pages):
        doc.add_heading(f"Page {i + 1}", level=2)
        for para in (txt.split("\n") if txt else [""]):
            doc.add_paragraph(para)
    doc.save(output_path)


# ════════════════════════════════════════════════════════════════════
#  CLI
# ════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import json as _json
    import sys
    args = sys.argv[1:]
    if not args or args[0] in ("-h", "--help"):
        print("usage:\n"
              "  python pdf_translate.py --status\n"
              "  python pdf_translate.py --text 'Hello' --to de [--from en]\n"
              "  python pdf_translate.py --image photo.png --to en [--from de]\n"
              "  python pdf_translate.py --pdf in.pdf out.txt --to en [--from auto]")
        raise SystemExit(0)

    def _opt(name, default=None):
        return args[args.index(name) + 1] if name in args else default

    if "--status" in args:
        print(_json.dumps(translation_status(), indent=2, ensure_ascii=False))
    elif "--text" in args:
        print(translate_text(_opt("--text"), _opt("--to"),
                             _opt("--from", "auto"))["translated"])
    elif "--image" in args:
        print(_json.dumps(translate_image(_opt("--image"), _opt("--to"),
                          _opt("--from", "auto")), indent=2, ensure_ascii=False))
    elif "--pdf" in args:
        i = args.index("--pdf")
        print(_json.dumps(translate_pdf(args[i + 1], args[i + 2], _opt("--to"),
                          _opt("--from", "auto")), indent=2, ensure_ascii=False))
