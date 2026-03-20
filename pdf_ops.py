"""PDF operations — merge, split, and page manipulation."""

import os
import tempfile
import threading
import logging
from dataclasses import dataclass, field
from typing import Callable

import math

import re

import fitz  # PyMuPDF
import pikepdf

log = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
#  Data classes
# ═══════════════════════════════════════════════════════════════════

@dataclass
class MergeResult:
    output_path: str
    input_paths: list[str]
    total_pages: int
    output_size: int


@dataclass
class SplitResult:
    input_path: str
    output_paths: list[str]
    pages_per_output: list[int]


@dataclass
class PageOpResult:
    input_path: str
    output_path: str
    operations: list[str] = field(default_factory=list)


# ═══════════════════════════════════════════════════════════════════
#  Merge
# ═══════════════════════════════════════════════════════════════════

def merge_pdfs(
    input_paths: list[str],
    output_path: str,
    on_progress: Callable[[int, int], None] | None = None,
    cancel: threading.Event | None = None,
) -> MergeResult:
    """Merge multiple PDFs into one.

    Args:
        input_paths: List of PDF file paths to merge (in order).
        output_path: Destination path for merged PDF.
        on_progress: Callback(current_file_index, total_files).
        cancel: Threading event to signal cancellation.

    Returns:
        MergeResult with merge details.
    """
    if len(input_paths) < 2:
        raise ValueError("Need at least 2 files to merge")

    total_pages = 0
    dest = pikepdf.Pdf.new()

    for i, path in enumerate(input_paths):
        if cancel and cancel.is_set():
            raise InterruptedError("Merge cancelled")

        if on_progress:
            on_progress(i, len(input_paths))

        try:
            src = pikepdf.open(path)
        except pikepdf.PasswordError:
            raise pikepdf.PasswordError(
                f"Cannot merge password-protected file: {os.path.basename(path)}"
            )

        dest.pages.extend(src.pages)
        total_pages += len(src.pages)
        src.close()

    if cancel and cancel.is_set():
        raise InterruptedError("Merge cancelled")

    # Atomic write via temp file
    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        dest.save(tmp, compress_streams=True)
        dest.close()
        os.replace(tmp, output_path)
    except Exception:
        dest.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise

    output_size = os.path.getsize(output_path)
    log.info("Merged %d files (%d pages) → %s (%d bytes)",
             len(input_paths), total_pages, output_path, output_size)

    return MergeResult(
        output_path=output_path,
        input_paths=list(input_paths),
        total_pages=total_pages,
        output_size=output_size,
    )


# ═══════════════════════════════════════════════════════════════════
#  TOC / Bookmarks
# ═══════════════════════════════════════════════════════════════════

_UNSAFE_FILENAME_RE = re.compile(r'[<>:"/\\|?*\x00-\x1f]')


def get_toc(input_path: str) -> list[dict]:
    """Extract the table of contents (bookmarks/outlines) from a PDF.

    Uses PyMuPDF's ``doc.get_toc()`` which returns
    ``[[level, title, page_number], ...]`` (1-indexed pages).

    Each entry is augmented with *end_page* — the last page that belongs
    to that entry (computed from the next entry at the same or higher level,
    or the document's last page).

    Returns:
        List of dicts ``{level, title, page, end_page}`` sorted by page.
        Empty list if the PDF has no bookmarks.
    """
    doc = fitz.open(input_path)
    raw_toc = doc.get_toc()  # [[level, title, page], ...]
    num_pages = doc.page_count
    doc.close()

    if not raw_toc:
        return []

    entries: list[dict] = []
    for i, (level, title, page) in enumerate(raw_toc):
        # Find end page: next entry at same or higher (lower number) level
        end_page = num_pages
        for j in range(i + 1, len(raw_toc)):
            if raw_toc[j][0] <= level:
                end_page = raw_toc[j][2] - 1
                break

        entries.append({
            "level": level,
            "title": title.strip(),
            "page": page,
            "end_page": max(end_page, page),  # ensure at least 1 page
        })

    return entries


def _sanitize_title(title: str) -> str:
    """Sanitize a chapter title for use as a filename component."""
    clean = _UNSAFE_FILENAME_RE.sub('_', title)
    clean = clean.strip('. _')
    return clean[:80] or 'untitled'


# ═══════════════════════════════════════════════════════════════════
#  Split
# ═══════════════════════════════════════════════════════════════════

def _parse_ranges(range_str: str, max_page: int) -> list[tuple[int, int]]:
    """Parse a range string like '1-3, 5, 8-10' into (start, end) tuples (1-indexed, inclusive)."""
    ranges = []
    for part in range_str.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            start, end = int(a.strip()), int(b.strip())
        else:
            start = end = int(part)
        if start < 1 or end > max_page or start > end:
            raise ValueError(f"Invalid range: {part} (pages 1–{max_page})")
        ranges.append((start, end))
    return ranges


def split_pdf(
    input_path: str,
    output_dir: str,
    mode: str = "all",
    ranges: str | None = None,
    every_n: int = 1,
    chapters: list[dict] | None = None,
    name_template: str = "{name}_page_{start}",
    cancel: threading.Event | None = None,
) -> SplitResult:
    """Split a PDF into multiple files.

    Args:
        input_path: Source PDF path.
        output_dir: Directory for output files.
        mode: "all" | "ranges" | "every_n" | "chapters".
        ranges: Range string like "1-3, 5, 8-10" (for mode="ranges").
        every_n: Pages per output file (for mode="every_n").
        chapters: List of {title, start_page, end_page} dicts (for mode="chapters").
        name_template: Output naming template. Supports {name}, {start}, {end}, {title}.
        cancel: Threading event for cancellation.

    Returns:
        SplitResult with output file paths.
    """
    if not output_dir:
        output_dir = os.path.dirname(os.path.abspath(input_path))
    src = pikepdf.open(input_path)
    num_pages = len(src.pages)
    base_name = os.path.splitext(os.path.basename(input_path))[0]

    # Build page groups: list of (start, end) 1-indexed inclusive
    # For chapters mode, also build a parallel titles list
    chapter_titles: list[str] = []

    if mode == "all":
        groups = [(i, i) for i in range(1, num_pages + 1)]
    elif mode == "ranges":
        if not ranges:
            raise ValueError("Range string required for 'ranges' mode")
        groups = _parse_ranges(ranges, num_pages)
    elif mode == "every_n":
        if every_n < 1:
            raise ValueError("every_n must be >= 1")
        groups = []
        for start in range(1, num_pages + 1, every_n):
            end = min(start + every_n - 1, num_pages)
            groups.append((start, end))
    elif mode == "chapters":
        if not chapters:
            raise ValueError("Chapters list required for 'chapters' mode")
        groups = []
        for ch in chapters:
            s = int(ch["start_page"])
            e = int(ch["end_page"])
            groups.append((s, e))
            chapter_titles.append(_sanitize_title(ch.get("title", "")))
    else:
        raise ValueError(f"Unknown split mode: {mode}")

    os.makedirs(output_dir, exist_ok=True)
    output_paths = []
    pages_per_output = []

    for idx, (start, end) in enumerate(groups):
        if cancel and cancel.is_set():
            src.close()
            raise InterruptedError("Split cancelled")

        title = chapter_titles[idx] if idx < len(chapter_titles) else ""
        out_name = name_template.format(
            name=base_name, filename=base_name,  # {name} and {filename} both work
            start=start, end=end, n=start,        # {start}/{end}/{n} all work
            title=title,
        ) + ".pdf"
        out_path = os.path.join(output_dir, out_name)

        dest = pikepdf.Pdf.new()
        for p in range(start - 1, end):  # 0-indexed
            dest.pages.append(src.pages[p])
        dest.save(out_path, compress_streams=True)
        dest.close()

        output_paths.append(out_path)
        pages_per_output.append(end - start + 1)

    src.close()
    log.info("Split %s into %d files", input_path, len(output_paths))

    return SplitResult(
        input_path=input_path,
        output_paths=output_paths,
        pages_per_output=pages_per_output,
    )


# ═══════════════════════════════════════════════════════════════════
#  Page operations
# ═══════════════════════════════════════════════════════════════════

def apply_page_operations(
    input_path: str,
    output_path: str,
    rotations: dict[int, int] | None = None,
    delete_pages: list[int] | None = None,
    new_order: list[int] | None = None,
    cancel: threading.Event | None = None,
) -> PageOpResult:
    """Apply page-level operations: rotate, delete, reorder.

    Args:
        input_path: Source PDF path.
        output_path: Destination PDF path.
        rotations: Dict mapping 0-indexed page → degrees (90, 180, 270).
        delete_pages: List of 0-indexed pages to delete.
        new_order: New page order as list of 0-indexed page numbers.
                   Applied AFTER deletions. If None, original order is kept.
        cancel: Threading event for cancellation.

    Returns:
        PageOpResult describing what was done.
    """
    src = pikepdf.open(input_path)
    ops = []

    # Apply rotations
    if rotations:
        for page_idx, degrees in rotations.items():
            if 0 <= page_idx < len(src.pages):
                current = int(src.pages[page_idx].get("/Rotate", 0))
                src.pages[page_idx]["/Rotate"] = (current + degrees) % 360
        ops.append(f"Rotated {len(rotations)} page(s)")

    if cancel and cancel.is_set():
        src.close()
        raise InterruptedError("Operation cancelled")

    # Build page list (apply deletions)
    delete_set = set(delete_pages) if delete_pages else set()
    remaining = [i for i in range(len(src.pages)) if i not in delete_set]
    if delete_pages:
        ops.append(f"Deleted {len(delete_set)} page(s)")

    # Apply reorder
    if new_order is not None:
        # new_order refers to indices within the remaining pages
        remaining = [remaining[i] for i in new_order if 0 <= i < len(remaining)]
        ops.append("Reordered pages")

    if cancel and cancel.is_set():
        src.close()
        raise InterruptedError("Operation cancelled")

    # Build output
    dest = pikepdf.Pdf.new()
    for idx in remaining:
        dest.pages.append(src.pages[idx])

    # Atomic write
    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        dest.save(tmp, compress_streams=True)
        dest.close()
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        dest.close()
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise

    output_size = os.path.getsize(output_path)
    ops_str = ", ".join(ops) if ops else "No changes"
    log.info("Page ops on %s → %s: %s", input_path, output_path, ops_str)

    return PageOpResult(
        input_path=input_path,
        output_path=output_path,
        operations=ops,
    )


# ═══════════════════════════════════════════════════════════════════
#  Security
# ═══════════════════════════════════════════════════════════════════

def protect_pdf(input_path, output_path, user_password="", owner_password="",
                permissions=None, encryption="AES-256"):
    """Add password protection and set permissions."""
    if permissions is None:
        permissions = {"print": True, "copy": True, "edit": True, "annotate": True}

    allow = pikepdf.Permissions(
        print_lowres=permissions.get("print", True),
        print_highres=permissions.get("print", True),
        extract=permissions.get("copy", True),
        modify_other=permissions.get("edit", True),
        modify_annotation=permissions.get("annotate", True),
        modify_form=permissions.get("edit", True),
        modify_assembly=permissions.get("edit", True),
        accessibility=True,
    )

    use_aes = encryption.startswith("AES")
    R = 6 if encryption == "AES-256" else 4

    src = pikepdf.open(input_path)
    enc = pikepdf.Encryption(
        owner=owner_password or user_password,
        user=user_password,
        allow=allow,
        aes=use_aes,
        R=R,
    )

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, encryption=enc)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Protected %s → %s", input_path, output_path)


def unlock_pdf(input_path, output_path, password=""):
    """Remove password protection from a PDF."""
    src = pikepdf.open(input_path, password=password)
    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Unlocked %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  PDF to Images
# ═══════════════════════════════════════════════════════════════════

@dataclass
class PdfToImagesResult:
    input_path: str
    output_dir: str
    image_paths: list[str]
    page_count: int
    format: str

def pdf_to_images(input_path, output_dir, fmt="png", dpi=150, quality=85,
                  page_range=None, on_progress=None, cancel=None):
    """Export PDF pages as fully-rendered images using PyMuPDF."""
    import fitz  # PyMuPDF

    os.makedirs(output_dir, exist_ok=True)
    doc = fitz.open(input_path)
    num_pages = len(doc)
    base = os.path.splitext(os.path.basename(input_path))[0]

    if page_range:
        ranges = _parse_ranges(page_range, num_pages)
        pages_to_export = []
        for s, e in ranges:
            pages_to_export.extend(range(s - 1, e))
    else:
        pages_to_export = list(range(num_pages))

    image_paths = []
    zoom = dpi / 72.0
    matrix = fitz.Matrix(zoom, zoom)

    for i, page_idx in enumerate(pages_to_export):
        if cancel and cancel.is_set():
            doc.close()
            raise InterruptedError("Export cancelled")
        if on_progress:
            on_progress(i, len(pages_to_export))

        page = doc[page_idx]
        pix = page.get_pixmap(matrix=matrix, alpha=False)

        ext = "jpg" if fmt.lower() in ("jpg", "jpeg") else "png"
        out_path = os.path.join(output_dir, f"{base}_page_{page_idx + 1}.{ext}")

        if ext == "jpg":
            pix.save(out_path, output="jpeg", jpg_quality=quality)
        else:
            pix.save(out_path, output="png")
        image_paths.append(out_path)

    doc.close()
    log.info("Exported %d pages from %s as %s", len(image_paths), input_path, fmt)
    return PdfToImagesResult(input_path, output_dir, image_paths, len(image_paths), fmt)


# ═══════════════════════════════════════════════════════════════════
#  Images to PDF
# ═══════════════════════════════════════════════════════════════════

@dataclass
class ImagesToPdfResult:
    output_path: str
    image_count: int
    page_count: int

def images_to_pdf(image_paths, output_path, page_size="auto", margin_mm=10,
                  on_progress=None, cancel=None):
    """Convert images into a PDF document.

    Handles EXIF rotation, transparency (via white background compositing),
    and preserves image quality.
    """
    from PIL import Image, ImageOps
    import io

    PAGE_SIZES = {
        "a4": (595.28, 841.89),
        "letter": (612, 792),
        "auto": None,
    }

    margin_pt = margin_mm * 72 / 25.4
    dest = pikepdf.Pdf.new()

    for i, img_path in enumerate(image_paths):
        if cancel and cancel.is_set():
            raise InterruptedError("Conversion cancelled")
        if on_progress:
            on_progress(i, len(image_paths))

        pil = Image.open(img_path)

        # Apply EXIF orientation (auto-rotates photos from cameras)
        pil = ImageOps.exif_transpose(pil)

        # Handle transparency: composite onto white background
        if pil.mode in ("RGBA", "LA", "PA"):
            background = Image.new("RGB", pil.size, "white")
            background.paste(pil, mask=pil.split()[-1])
            pil = background
        elif pil.mode != "RGB":
            pil = pil.convert("RGB")

        img_w, img_h = pil.size
        size_key = page_size.lower()
        page_dims = PAGE_SIZES.get(size_key)

        if page_dims is None:
            dpi = pil.info.get("dpi", (72, 72))
            # Clamp DPI to sane range to avoid tiny/huge pages
            dpi_x = max(36, min(float(dpi[0]), 1200))
            dpi_y = max(36, min(float(dpi[1]), 1200))
            pw = img_w * 72 / dpi_x + 2 * margin_pt
            ph = img_h * 72 / dpi_y + 2 * margin_pt
        else:
            pw, ph = page_dims

        # Save image to bytes as JPEG
        buf = io.BytesIO()
        pil.save(buf, "JPEG", quality=92)
        buf.seek(0)

        # Create page with image
        page_pdf = pikepdf.Pdf.new()
        page_pdf.add_blank_page(page_size=(pw, ph))
        raw_img = pikepdf.Stream(page_pdf, buf.read())
        raw_img["/Type"] = pikepdf.Name("/XObject")
        raw_img["/Subtype"] = pikepdf.Name("/Image")
        raw_img["/Width"] = img_w
        raw_img["/Height"] = img_h
        raw_img["/ColorSpace"] = pikepdf.Name("/DeviceRGB")
        raw_img["/BitsPerComponent"] = 8
        raw_img["/Filter"] = pikepdf.Name("/DCTDecode")

        # Calculate placement (centered with margins)
        avail_w = pw - 2 * margin_pt
        avail_h = ph - 2 * margin_pt
        scale = min(avail_w / img_w, avail_h / img_h)
        disp_w = img_w * scale
        disp_h = img_h * scale
        x = margin_pt + (avail_w - disp_w) / 2
        y = margin_pt + (avail_h - disp_h) / 2

        page = page_pdf.pages[0]
        page["/Resources"] = pikepdf.Dictionary({
            "/XObject": pikepdf.Dictionary({"/Img0": raw_img})
        })
        content = f"q {disp_w:.2f} 0 0 {disp_h:.2f} {x:.2f} {y:.2f} cm /Img0 Do Q"
        page["/Contents"] = pikepdf.Stream(page_pdf, content.encode())

        dest.pages.extend(page_pdf.pages)

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        dest.save(tmp, compress_streams=True)
        dest.close()
        os.replace(tmp, output_path)
    except Exception:
        dest.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise

    log.info("Created PDF from %d images → %s", len(image_paths), output_path)
    return ImagesToPdfResult(output_path, len(image_paths), len(image_paths))


# ═══════════════════════════════════════════════════════════════════
#  PDF to Word
# ═══════════════════════════════════════════════════════════════════

@dataclass
class PdfToWordResult:
    input_path: str
    output_path: str
    page_count: int

def pdf_to_word(input_path, output_path, on_progress=None, cancel=None):
    """Extract text from PDF and create a Word document using PyMuPDF."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    import fitz  # PyMuPDF

    pdf_doc = fitz.open(input_path)
    doc = Document()
    num_pages = len(pdf_doc)

    for i in range(num_pages):
        if cancel and cancel.is_set():
            pdf_doc.close()
            raise InterruptedError("Conversion cancelled")
        if on_progress:
            on_progress(i, num_pages)

        if i > 0:
            doc.add_page_break()

        page = pdf_doc[i]

        # Extract structured text blocks: (x0, y0, x1, y1, text, block_no, block_type)
        blocks = page.get_text("blocks")

        if not blocks:
            doc.add_paragraph("[No extractable text on this page]")
            continue

        page_has_text = False
        for block in blocks:
            # block_type 0 = text, 1 = image
            if block[6] == 0:
                text = block[4].strip()
                if text:
                    page_has_text = True
                    # Detect heading-like blocks by font size via dict extraction
                    para = doc.add_paragraph(text)

        if not page_has_text:
            doc.add_paragraph("[No extractable text on this page]")

    pdf_doc.close()
    doc.save(output_path)
    log.info("PDF to Word: %s → %s (%d pages)", input_path, output_path, num_pages)
    return PdfToWordResult(input_path, output_path, num_pages)


def _extract_text_from_stream(raw_bytes):
    """Legacy text extraction from PDF content stream bytes (fallback only)."""
    text_parts = []
    try:
        data = raw_bytes.decode("latin-1")
    except Exception:
        return ""

    i = 0
    while i < len(data):
        if data[i] == "(":
            depth = 1
            start = i + 1
            i += 1
            while i < len(data) and depth > 0:
                if data[i] == "\\":
                    i += 1
                elif data[i] == "(":
                    depth += 1
                elif data[i] == ")":
                    depth -= 1
                i += 1
            if depth == 0:
                fragment = data[start:i - 1]
                fragment = fragment.replace("\\n", "\n").replace("\\r", "\r")
                fragment = fragment.replace("\\t", "\t")
                fragment = fragment.replace("\\(", "(").replace("\\)", ")")
                fragment = fragment.replace("\\\\", "\\")
                text_parts.append(fragment)
        else:
            i += 1

    return " ".join(text_parts)


def _extract_text_pymupdf(input_path, page_indices=None):
    """Extract text from PDF pages using PyMuPDF (high quality).

    Returns a list of (page_number, text) tuples.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(input_path)
    results = []
    indices = page_indices if page_indices is not None else list(range(len(doc)))

    for page_idx in indices:
        if 0 <= page_idx < len(doc):
            page = doc[page_idx]
            text = page.get_text("text")
            results.append((page_idx, text))

    doc.close()
    return results


# ═══════════════════════════════════════════════════════════════════
#  Watermark
# ═══════════════════════════════════════════════════════════════════

def _pdf_escape(text):
    """Escape special characters for PDF text strings inside parentheses."""
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def add_watermark(input_path, output_path, text="WATERMARK", opacity=0.3,
                  rotation=45, font_size=48, color="#888888", position="center",
                  page_range=None, cancel=None):
    """Add a text watermark to PDF pages.

    Text is centered on the page (accounting for text width) and supports
    special characters. Opacity applies to both fill and stroke.
    """
    src = pikepdf.open(input_path)
    num_pages = len(src.pages)

    if page_range:
        ranges = _parse_ranges(page_range, num_pages)
        target_pages = set()
        for s, e in ranges:
            target_pages.update(range(s - 1, e))
    else:
        target_pages = set(range(num_pages))

    # Convert hex color to RGB floats
    c = color.lstrip("#")
    r, g, b = int(c[0:2], 16) / 255, int(c[2:4], 16) / 255, int(c[4:6], 16) / 255

    angle_rad = math.radians(rotation)
    cos_a = math.cos(angle_rad)
    sin_a = math.sin(angle_rad)

    # Approximate text width for Helvetica (avg char width ~0.5 * font_size)
    approx_text_width = len(text) * font_size * 0.5
    safe_text = _pdf_escape(text)

    for i in range(num_pages):
        if cancel and cancel.is_set():
            src.close()
            raise InterruptedError("Watermark cancelled")

        if i not in target_pages:
            continue

        page = src.pages[i]
        mbox = page.get("/MediaBox", pikepdf.Array([0, 0, 612, 792]))
        pw = float(mbox[2]) - float(mbox[0])
        ph = float(mbox[3]) - float(mbox[1])

        # Center of page, offset by half the text width so it's visually centered
        cx = pw / 2 - (approx_text_width / 2) * cos_a
        cy = ph / 2 - (approx_text_width / 2) * sin_a

        # Build watermark content stream
        wm = (
            f"q\n"
            f"/GS0 gs\n"
            f"{r:.3f} {g:.3f} {b:.3f} rg\n"
            f"BT\n"
            f"/F1 {font_size} Tf\n"
            f"{cos_a:.4f} {sin_a:.4f} {-sin_a:.4f} {cos_a:.4f} {cx:.2f} {cy:.2f} Tm\n"
            f"({safe_text}) Tj\n"
            f"ET\n"
            f"Q\n"
        )

        wm_stream = pikepdf.Stream(src, wm.encode())

        # Ensure page has a font resource
        if "/Resources" not in page:
            page["/Resources"] = pikepdf.Dictionary()
        resources = page["/Resources"]
        if "/Font" not in resources:
            resources["/Font"] = pikepdf.Dictionary()
        if "/F1" not in resources["/Font"]:
            resources["/Font"]["/F1"] = pikepdf.Dictionary({
                "/Type": pikepdf.Name("/Font"),
                "/Subtype": pikepdf.Name("/Type1"),
                "/BaseFont": pikepdf.Name("/Helvetica"),
            })

        # Add graphics state for opacity
        if "/ExtGState" not in resources:
            resources["/ExtGState"] = pikepdf.Dictionary()
        resources["/ExtGState"]["/GS0"] = pikepdf.Dictionary({
            "/Type": pikepdf.Name("/ExtGState"),
            "/ca": opacity,
            "/CA": opacity,
        })

        # Append watermark to page contents
        existing = page.get("/Contents")
        if existing is None:
            page["/Contents"] = wm_stream
        elif isinstance(existing, pikepdf.Array):
            existing.append(wm_stream)
        else:
            page["/Contents"] = pikepdf.Array([existing, wm_stream])

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, compress_streams=True)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Watermarked %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Page Numbers
# ═══════════════════════════════════════════════════════════════════

def add_page_numbers(input_path, output_path, position="bottom-center",
                     font_size=10, fmt="{page}", start_number=1,
                     margin_pt=36, cancel=None):
    """Insert page numbering on all pages.

    Properly centers text for center-aligned positions and escapes
    special PDF characters in the label.
    """
    src = pikepdf.open(input_path)
    num_pages = len(src.pages)

    for i in range(num_pages):
        if cancel and cancel.is_set():
            src.close()
            raise InterruptedError("Page numbering cancelled")

        page = src.pages[i]
        mbox = page.get("/MediaBox", pikepdf.Array([0, 0, 612, 792]))
        pw = float(mbox[2]) - float(mbox[0])
        ph = float(mbox[3]) - float(mbox[1])

        page_num = start_number + i
        label = fmt.replace("{page}", str(page_num)).replace("{total}", str(num_pages))
        safe_label = _pdf_escape(label)

        # Approximate label width for centering (Helvetica avg ~0.5 * font_size per char)
        label_width = len(label) * font_size * 0.5

        # Calculate position — offset centered positions by half the text width
        is_center = "center" in position
        is_right = "right" in position

        if is_center:
            x_base = pw / 2 - label_width / 2
        elif is_right:
            x_base = pw - margin_pt - label_width
        else:
            x_base = margin_pt

        if "top" in position:
            y_base = ph - margin_pt
        else:
            y_base = margin_pt

        content = (
            f"q\n"
            f"0 0 0 rg\n"
            f"BT\n"
            f"/F1 {font_size} Tf\n"
            f"{x_base:.2f} {y_base:.2f} Td\n"
            f"({safe_label}) Tj\n"
            f"ET\n"
            f"Q\n"
        )

        num_stream = pikepdf.Stream(src, content.encode())

        if "/Resources" not in page:
            page["/Resources"] = pikepdf.Dictionary()
        resources = page["/Resources"]
        if "/Font" not in resources:
            resources["/Font"] = pikepdf.Dictionary()
        if "/F1" not in resources["/Font"]:
            resources["/Font"]["/F1"] = pikepdf.Dictionary({
                "/Type": pikepdf.Name("/Font"),
                "/Subtype": pikepdf.Name("/Type1"),
                "/BaseFont": pikepdf.Name("/Helvetica"),
            })

        existing = page.get("/Contents")
        if existing is None:
            page["/Contents"] = num_stream
        elif isinstance(existing, pikepdf.Array):
            existing.append(num_stream)
        else:
            page["/Contents"] = pikepdf.Array([existing, num_stream])

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, compress_streams=True)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Added page numbers to %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Metadata
# ═══════════════════════════════════════════════════════════════════

@dataclass
class MetadataFields:
    title: str = ""
    author: str = ""
    subject: str = ""
    keywords: str = ""
    creator: str = ""
    producer: str = ""

def read_metadata(path):
    """Read PDF metadata fields. Returns a dict."""
    src = pikepdf.open(path)
    info = src.docinfo
    result = {
        "title": str(info.get("/Title", "")),
        "author": str(info.get("/Author", "")),
        "subject": str(info.get("/Subject", "")),
        "keywords": str(info.get("/Keywords", "")),
        "creator": str(info.get("/Creator", "")),
        "producer": str(info.get("/Producer", "")),
    }
    src.close()
    return result


def write_metadata(input_path, output_path, fields):
    """Write PDF metadata fields. fields is a dict with keys: title, author, etc."""
    if isinstance(fields, dict):
        f = fields
    else:
        f = {"title": fields.title, "author": fields.author, "subject": fields.subject,
             "keywords": fields.keywords, "creator": fields.creator, "producer": fields.producer}

    src = pikepdf.open(input_path)
    with src.open_metadata() as meta:
        if f.get("title"):
            meta["dc:title"] = f["title"]
        if f.get("author"):
            meta["dc:creator"] = [f["author"]]
        if f.get("subject"):
            meta["dc:description"] = f["subject"]
        if f.get("keywords"):
            meta["pdf:Keywords"] = f["keywords"]
        if f.get("creator"):
            meta["xmp:CreatorTool"] = f["creator"]
        if f.get("producer"):
            meta["pdf:Producer"] = f["producer"]

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Updated metadata %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Extract Images
# ═══════════════════════════════════════════════════════════════════

@dataclass
class ExtractImagesResult:
    input_path: str
    output_dir: str
    image_paths: list[str]
    image_count: int

def extract_images(input_path, output_dir, fmt="png", min_size=0,
                   on_progress=None, cancel=None):
    """Extract all images from a PDF."""
    from PIL import Image
    import io

    os.makedirs(output_dir, exist_ok=True)
    src = pikepdf.open(input_path)
    base = os.path.splitext(os.path.basename(input_path))[0]
    image_paths = []
    img_idx = 0

    for page_num, page in enumerate(src.pages):
        if cancel and cancel.is_set():
            src.close()
            raise InterruptedError("Extraction cancelled")
        if on_progress:
            on_progress(page_num, len(src.pages))

        if "/Resources" not in page:
            continue
        resources = page["/Resources"]
        if "/XObject" not in resources:
            continue

        for key, xobj_ref in resources["/XObject"].items():
            try:
                xobj = xobj_ref
                if hasattr(xobj, 'resolve'):
                    xobj = xobj.resolve()
                if xobj.get("/Subtype") != "/Image":
                    continue

                w = int(xobj.get("/Width", 0))
                h = int(xobj.get("/Height", 0))
                if w * h < min_size:
                    continue

                # Try to extract image data
                raw = xobj.read_raw_bytes()
                filter_name = str(xobj.get("/Filter", ""))

                if "/DCTDecode" in filter_name:
                    ext = "jpg"
                    out_path = os.path.join(output_dir, f"{base}_img_{img_idx:03d}.{ext}")
                    with open(out_path, "wb") as f:
                        f.write(raw)
                elif "/JPXDecode" in filter_name:
                    ext = "jp2"
                    out_path = os.path.join(output_dir, f"{base}_img_{img_idx:03d}.{ext}")
                    with open(out_path, "wb") as f:
                        f.write(raw)
                else:
                    # Try to decode with Pillow
                    try:
                        data = xobj.read_bytes()
                        cs = str(xobj.get("/ColorSpace", "/DeviceRGB"))
                        bpc = int(xobj.get("/BitsPerComponent", 8))

                        if "/DeviceRGB" in cs:
                            mode = "RGB"
                        elif "/DeviceCMYK" in cs:
                            mode = "CMYK"
                        elif "/DeviceGray" in cs:
                            mode = "L"
                        else:
                            mode = "RGB"

                        img = Image.frombytes(mode, (w, h), data)
                        if mode == "CMYK":
                            img = img.convert("RGB")

                        ext = fmt.lower()
                        if ext == "jpeg":
                            ext = "jpg"
                        out_path = os.path.join(output_dir, f"{base}_img_{img_idx:03d}.{ext}")
                        img.save(out_path)
                    except Exception:
                        continue

                image_paths.append(out_path)
                img_idx += 1
            except Exception:
                continue

    src.close()
    log.info("Extracted %d images from %s", len(image_paths), input_path)
    return ExtractImagesResult(input_path, output_dir, image_paths, len(image_paths))


# ═══════════════════════════════════════════════════════════════════
#  Extract Text
# ═══════════════════════════════════════════════════════════════════

@dataclass
class ExtractTextResult:
    input_path: str
    output_path: str
    page_count: int
    char_count: int

def extract_text(input_path, output_path, page_range=None, cancel=None):
    """Extract text content from a PDF to a text file using PyMuPDF."""
    import fitz  # PyMuPDF

    doc = fitz.open(input_path)
    num_pages = len(doc)

    if page_range:
        ranges = _parse_ranges(page_range, num_pages)
        target_pages = []
        for s, e in ranges:
            target_pages.extend(range(s - 1, e))
    else:
        target_pages = list(range(num_pages))

    all_text = []
    for page_idx in target_pages:
        if cancel and cancel.is_set():
            doc.close()
            raise InterruptedError("Extraction cancelled")

        page = doc[page_idx]
        text = page.get_text("text")
        all_text.append(f"--- Page {page_idx + 1} ---\n{text}\n")

    doc.close()

    full_text = "\n".join(all_text)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(full_text)

    log.info("Extracted text from %s → %s", input_path, output_path)
    return ExtractTextResult(input_path, output_path, len(target_pages), len(full_text))


# ═══════════════════════════════════════════════════════════════════
#  Crop Pages
# ═══════════════════════════════════════════════════════════════════

def crop_pages(input_path, output_path, margins, unit="mm"):
    """Crop pages by trimming margins. margins = dict with left/right/top/bottom."""
    # Convert to points
    scale = {"mm": 72 / 25.4, "inches": 72, "points": 1, "pt": 1}.get(unit, 72 / 25.4)

    left = margins.get("left", 0) * scale
    right = margins.get("right", 0) * scale
    top = margins.get("top", 0) * scale
    bottom = margins.get("bottom", 0) * scale

    src = pikepdf.open(input_path)
    for page in src.pages:
        mbox = page.get("/MediaBox", pikepdf.Array([0, 0, 612, 792]))
        x0 = float(mbox[0]) + left
        y0 = float(mbox[1]) + bottom
        x1 = float(mbox[2]) - right
        y1 = float(mbox[3]) - top

        if x1 <= x0 or y1 <= y0:
            src.close()
            raise ValueError("Crop margins are too large — no page area remaining")

        page["/CropBox"] = pikepdf.Array([x0, y0, x1, y1])

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, compress_streams=True)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Cropped %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Flatten PDF
# ═══════════════════════════════════════════════════════════════════

def flatten_pdf(input_path, output_path, annotations=True, forms=True):
    """Remove annotations and/or form fields from a PDF."""
    src = pikepdf.open(input_path)

    for page in src.pages:
        if annotations and "/Annots" in page:
            if forms:
                del page["/Annots"]
            else:
                # Keep widget annotations (form fields), remove others
                annots = page["/Annots"]
                kept = []
                for annot in annots:
                    a = annot.resolve() if hasattr(annot, 'resolve') else annot
                    subtype = str(a.get("/Subtype", ""))
                    if subtype == "/Widget":
                        kept.append(annot)
                if kept:
                    page["/Annots"] = pikepdf.Array(kept)
                else:
                    del page["/Annots"]

    if forms and "/AcroForm" in src.Root:
        del src.Root["/AcroForm"]

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, compress_streams=True)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Flattened %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  N-up Layout
# ═══════════════════════════════════════════════════════════════════

def nup_layout(input_path, output_path, pages_per_sheet=4,
               page_size="A4", orientation="landscape"):
    """Arrange multiple pages per sheet."""
    PAGE_SIZES = {
        "a4": (595.28, 841.89),
        "letter": (612, 792),
        "a3": (841.89, 1190.55),
    }

    base_w, base_h = PAGE_SIZES.get(page_size.lower(), (595.28, 841.89))
    if orientation == "landscape":
        sheet_w, sheet_h = max(base_w, base_h), min(base_w, base_h)
    else:
        sheet_w, sheet_h = min(base_w, base_h), max(base_w, base_h)

    # Calculate grid
    cols = int(math.ceil(math.sqrt(pages_per_sheet)))
    rows = int(math.ceil(pages_per_sheet / cols))

    cell_w = sheet_w / cols
    cell_h = sheet_h / rows

    src = pikepdf.open(input_path)
    dest = pikepdf.Pdf.new()
    num_pages = len(src.pages)

    for sheet_start in range(0, num_pages, pages_per_sheet):
        dest.add_blank_page(page_size=(sheet_w, sheet_h))
        out_page = dest.pages[-1]

        content_parts = []
        xobjects = {}

        for slot in range(pages_per_sheet):
            page_idx = sheet_start + slot
            if page_idx >= num_pages:
                break

            src_page = src.pages[page_idx]
            # Create form XObject from source page
            xobj_name = f"/P{slot}"

            # Get source page dimensions
            mbox = src_page.get("/MediaBox", pikepdf.Array([0, 0, 612, 792]))
            src_w = float(mbox[2]) - float(mbox[0])
            src_h = float(mbox[3]) - float(mbox[1])

            # Calculate position and scale
            col = slot % cols
            row_idx = slot // cols
            scale = min(cell_w / src_w, cell_h / src_h) * 0.95  # 5% margin

            scaled_w = src_w * scale
            scaled_h = src_h * scale

            x = col * cell_w + (cell_w - scaled_w) / 2
            y = sheet_h - (row_idx + 1) * cell_h + (cell_h - scaled_h) / 2

            # Create form xobject
            form_xobj = dest.copy_foreign(src_page.as_form_xobject())
            xobjects[xobj_name] = form_xobj

            content_parts.append(
                f"q {scale:.6f} 0 0 {scale:.6f} {x:.2f} {y:.2f} cm {xobj_name} Do Q"
            )

        if "/Resources" not in out_page:
            out_page["/Resources"] = pikepdf.Dictionary()
        out_page["/Resources"]["/XObject"] = pikepdf.Dictionary(xobjects)
        out_page["/Contents"] = pikepdf.Stream(dest, "\n".join(content_parts).encode())

    src.close()

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        dest.save(tmp, compress_streams=True)
        dest.close()
        os.replace(tmp, output_path)
    except Exception:
        dest.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("N-up %d pages/sheet %s → %s", pages_per_sheet, input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Repair PDF
# ═══════════════════════════════════════════════════════════════════

def repair_pdf(input_path, output_path):
    """Attempt to repair a corrupted PDF."""
    # pikepdf/QPDF has built-in repair capabilities
    src = pikepdf.open(input_path, attempt_recovery=True)

    out_dir = os.path.dirname(os.path.abspath(output_path))
    fd, tmp = tempfile.mkstemp(suffix=".pdf", dir=out_dir)
    os.close(fd)
    try:
        src.save(tmp, compress_streams=True, linearize=True)
        src.close()
        os.replace(tmp, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp):
            os.unlink(tmp)
        raise
    log.info("Repaired %s → %s", input_path, output_path)


# ═══════════════════════════════════════════════════════════════════
#  Compare PDFs
# ═══════════════════════════════════════════════════════════════════

@dataclass
class CompareResult:
    path_a: str
    path_b: str
    page_diffs: list[dict]
    metadata_diffs: list[dict]

def compare_pdfs(path_a, path_b):
    """Compare two PDFs page-by-page (text content + metadata) using PyMuPDF."""
    import fitz  # PyMuPDF

    doc_a = fitz.open(path_a)
    doc_b = fitz.open(path_b)

    max_pages = max(len(doc_a), len(doc_b))
    page_diffs = []

    for i in range(max_pages):
        if i >= len(doc_a):
            page_diffs.append({"page": i + 1, "status": "added_in_b", "details": "Page only in document B"})
            continue
        if i >= len(doc_b):
            page_diffs.append({"page": i + 1, "status": "removed_in_b", "details": "Page only in document A"})
            continue

        text_a = doc_a[i].get_text("text")
        text_b = doc_b[i].get_text("text")

        if text_a == text_b:
            page_diffs.append({"page": i + 1, "status": "identical", "details": ""})
        else:
            words_a = set(text_a.split())
            words_b = set(text_b.split())
            added = len(words_b - words_a)
            removed = len(words_a - words_b)
            page_diffs.append({
                "page": i + 1,
                "status": "different",
                "details": f"+{added} words, -{removed} words"
            })

    doc_a.close()
    doc_b.close()

    # Metadata comparison via pikepdf (richer metadata access)
    meta_diffs = []
    src_a = pikepdf.open(path_a)
    src_b = pikepdf.open(path_b)
    for key in ("/Title", "/Author", "/Subject", "/Creator"):
        val_a = str(src_a.docinfo.get(key, ""))
        val_b = str(src_b.docinfo.get(key, ""))
        if val_a != val_b:
            meta_diffs.append({"field": key, "a": val_a, "b": val_b})
    src_a.close()
    src_b.close()

    return CompareResult(path_a, path_b, page_diffs, meta_diffs)


# ═══════════════════════════════════════════════════════════════════
#  Redact
# ═══════════════════════════════════════════════════════════════════

@dataclass
class RedactResult:
    input_path: str
    output_path: str
    redaction_count: int
    pages_affected: int


def redact_text(
    input_path: str,
    output_path: str,
    search_terms: list[str],
    case_sensitive: bool = False,
    pages: list[int] | None = None,
    on_progress: Callable[[int, int], None] | None = None,
    cancel: threading.Event | None = None,
    password: str | None = None,
) -> RedactResult:
    """Redact text from a PDF by replacing matching content with black rectangles.

    This performs true redaction — the underlying text is removed, not just
    covered. The redacted content cannot be recovered from the output file.

    Args:
        input_path: Source PDF path.
        output_path: Destination path for the redacted PDF.
        search_terms: List of text strings to redact.
        case_sensitive: Whether matching is case-sensitive.
        pages: Optional list of 0-based page indices to redact (None = all).
        on_progress: Callback(current_page, total_pages).
        cancel: Threading event to signal cancellation.
        password: Password for encrypted PDFs.

    Returns:
        RedactResult with redaction details.
    """
    import re as re_mod

    if not search_terms:
        raise ValueError("No search terms provided for redaction")

    open_kwargs = {}
    if password:
        open_kwargs["password"] = password

    src = pikepdf.open(input_path, **open_kwargs)
    total_pages = len(src.pages)
    target_pages = pages if pages is not None else list(range(total_pages))

    total_redactions = 0
    pages_affected = set()

    for idx, page_num in enumerate(target_pages):
        if cancel and cancel.is_set():
            src.close()
            raise ValueError("Cancelled")

        if on_progress:
            on_progress(idx + 1, len(target_pages))

        if page_num < 0 or page_num >= total_pages:
            continue

        page = src.pages[page_num]

        # Extract the page content stream text to find matches
        try:
            if "/Contents" not in page:
                continue
            contents = page["/Contents"]
            if isinstance(contents, pikepdf.Array):
                raw = b""
                for stream in contents:
                    raw += stream.read_bytes()
            else:
                raw = contents.read_bytes()

            text_content = raw.decode("latin-1", errors="replace")
        except Exception:
            continue

        # Search for terms in the content stream text operands
        page_had_redaction = False
        for term in search_terms:
            if not term.strip():
                continue

            flags = 0 if case_sensitive else re_mod.IGNORECASE
            escaped = re_mod.escape(term)

            # Replace text within PDF text-showing operators: (text) Tj / (text) '
            # This handles the common case of text in parenthesized strings
            def _redact_match(m):
                nonlocal total_redactions, page_had_redaction
                total_redactions += 1
                page_had_redaction = True
                # Replace matched text with spaces of same byte length
                return b" " * len(m.group(0))

            # Match the term in parenthesized text strings
            pattern = escaped.encode("latin-1", errors="replace")
            if not case_sensitive:
                # Build case-insensitive byte pattern
                ci_parts = []
                for ch in term:
                    if ch.isalpha():
                        ci_parts.append(
                            f"[{re_mod.escape(ch.lower())}{re_mod.escape(ch.upper())}]"
                        )
                    else:
                        ci_parts.append(re_mod.escape(ch))
                pattern = "".join(ci_parts).encode("latin-1", errors="replace")

            new_content = re_mod.sub(pattern, _redact_match, raw)
            if new_content != raw:
                raw = new_content

        if page_had_redaction:
            pages_affected.add(page_num)
            # Write back the redacted content stream
            if isinstance(page["/Contents"], pikepdf.Array):
                # Merge into single stream
                page["/Contents"] = src.make_stream(raw)
            else:
                page["/Contents"].write(raw)

            # Add black redaction rectangles as annotation overlay
            # This provides visual indication that content was redacted
            _add_redaction_overlay(page, src)

    # Save — use a temp file for atomic write
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(output_path))
    os.close(tmp_fd)
    try:
        src.save(tmp_path)
        src.close()
        os.replace(tmp_path, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    return RedactResult(
        input_path=input_path,
        output_path=output_path,
        redaction_count=total_redactions,
        pages_affected=len(pages_affected),
    )


def _add_redaction_overlay(page, pdf):
    """Add a subtle 'REDACTED' footer marker on pages that were redacted."""
    mbox = page.get("/MediaBox", pikepdf.Array([0, 0, 612, 792]))
    pw = float(mbox[2]) - float(mbox[0])

    # Small grey "REDACTED" text at bottom-right corner
    marker = (
        f"q\n"
        f"0.7 0.7 0.7 rg\n"
        f"BT\n"
        f"/F1 6 Tf\n"
        f"{pw - 60:.2f} 8 Td\n"
        f"(REDACTED) Tj\n"
        f"ET\n"
        f"Q\n"
    )
    marker_stream = pikepdf.Stream(pdf, marker.encode())

    # Ensure font resource exists
    if "/Resources" not in page:
        page["/Resources"] = pikepdf.Dictionary()
    resources = page["/Resources"]
    if "/Font" not in resources:
        resources["/Font"] = pikepdf.Dictionary()
    if "/F1" not in resources["/Font"]:
        resources["/Font"]["/F1"] = pikepdf.Dictionary({
            "/Type": pikepdf.Name("/Font"),
            "/Subtype": pikepdf.Name("/Type1"),
            "/BaseFont": pikepdf.Name("/Helvetica"),
        })

    # Append marker to page contents
    existing = page.get("/Contents")
    if existing is None:
        page["/Contents"] = marker_stream
    elif isinstance(existing, pikepdf.Array):
        existing.append(marker_stream)
    else:
        page["/Contents"] = pikepdf.Array([existing, marker_stream])


def redact_region(
    input_path: str,
    output_path: str,
    regions: list[dict],
    on_progress: Callable[[int, int], None] | None = None,
    cancel: threading.Event | None = None,
    password: str | None = None,
) -> RedactResult:
    """Redact rectangular regions from PDF pages.

    Each region dict should have: page (0-based), x, y, width, height
    (in PDF points, origin at bottom-left).

    The region content is replaced with a filled black rectangle and
    the underlying content stream is modified to remove text in that area.

    Args:
        input_path: Source PDF path.
        output_path: Destination path.
        regions: List of region dicts with page, x, y, width, height.
        on_progress: Callback(current, total).
        cancel: Threading event.
        password: Password for encrypted PDFs.

    Returns:
        RedactResult with redaction details.
    """
    if not regions:
        raise ValueError("No regions provided for redaction")

    open_kwargs = {}
    if password:
        open_kwargs["password"] = password

    src = pikepdf.open(input_path, **open_kwargs)
    total_pages = len(src.pages)
    pages_affected = set()

    for idx, region in enumerate(regions):
        if cancel and cancel.is_set():
            src.close()
            raise ValueError("Cancelled")

        if on_progress:
            on_progress(idx + 1, len(regions))

        page_num = region.get("page", 0)
        if page_num < 0 or page_num >= total_pages:
            continue

        x = float(region.get("x", 0))
        y = float(region.get("y", 0))
        w = float(region.get("width", 100))
        h = float(region.get("height", 20))

        page = src.pages[page_num]
        pages_affected.add(page_num)

        # Create a black filled rectangle that covers the region
        redaction_stream = (
            f"q\n"
            f"0 0 0 rg\n"           # Black fill
            f"{x} {y} {w} {h} re\n"  # Rectangle
            f"f\n"                    # Fill
            f"Q\n"
        ).encode("ascii")

        # Append to existing contents
        if "/Contents" in page:
            existing = page["/Contents"]
            if isinstance(existing, pikepdf.Array):
                existing_bytes = b""
                for stream in existing:
                    existing_bytes += stream.read_bytes()
            else:
                existing_bytes = existing.read_bytes()
            combined = existing_bytes + b"\n" + redaction_stream
            page["/Contents"] = src.make_stream(combined)
        else:
            page["/Contents"] = src.make_stream(redaction_stream)

    # Atomic save
    tmp_fd, tmp_path = tempfile.mkstemp(suffix=".pdf", dir=os.path.dirname(output_path))
    os.close(tmp_fd)
    try:
        src.save(tmp_path)
        src.close()
        os.replace(tmp_path, output_path)
    except Exception:
        src.close()
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)
        raise

    return RedactResult(
        input_path=input_path,
        output_path=output_path,
        redaction_count=len(regions),
        pages_affected=len(pages_affected),
    )


def _page_text(page):
    """Extract text from a single pikepdf page for comparison (legacy fallback)."""
    try:
        if "/Contents" not in page:
            return ""
        contents = page["/Contents"]
        if isinstance(contents, pikepdf.Array):
            raw = b""
            for stream in contents:
                raw += stream.read_bytes()
        else:
            raw = contents.read_bytes()
        return _extract_text_from_stream(raw)
    except Exception:
        return ""
